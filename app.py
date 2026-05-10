import base64
import io
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from scorer import load_criteria, score_text

_APP_DIR = Path(__file__).resolve().parent

_ESCUDO_REMOTE_URL = (
    "https://raw.githubusercontent.com/claudiomlarrea/valorador_informes_finales/"
    "main/assets/escudo_uccuyo.png"
)


def _resolve_escudo_path() -> Path | None:
    assets = _APP_DIR / "assets"
    if not assets.is_dir():
        return None
    for name in ("escudo_uccuyo.png", "escudo_uccuyo.jpg", "escudo_uccuyo.jpeg"):
        p = assets / name
        if p.is_file():
            return p
    return None


def _escudo_src_for_banner() -> str:
    p = _resolve_escudo_path()
    if p is not None:
        ext = p.suffix.lower()
        mime = "image/jpeg" if ext in (".jpg", ".jpeg") else "image/png"
        b64 = base64.standard_b64encode(p.read_bytes()).decode("ascii")
        return f"data:{mime};base64,{b64}"
    return _ESCUDO_REMOTE_URL


_UCCI_GLOBAL_CSS = """
<style>
:root {
    --ucc-green: #00664d;
    --ucc-green-dark: #00523e;
    --ucc-accent: #28a745;
    --ucc-page-bg: #E6E6E6;
    --ucc-sidebar-bg: #262730;
    --ucc-text: #262730;
    --ucc-heading-card: #2c3838;
    --ucc-lead-muted: #5f6b6f;
}

.stApp {
    background-color: var(--ucc-page-bg);
}

header[data-testid="stHeader"] {
    background: var(--ucc-page-bg) !important;
    border-bottom: 1px solid rgba(0, 0, 0, 0.06);
}
div[data-testid="stDecoration"] {
    height: 3px !important;
    margin-top: env(safe-area-inset-top, 0);
    background: linear-gradient(
        90deg,
        var(--ucc-green-dark) 0%,
        var(--ucc-green) 50%,
        var(--ucc-green-dark) 100%
    ) !important;
}

.block-container {
    padding-top: 2rem !important;
    padding-left: calc(1rem + env(safe-area-inset-left, 0px)) !important;
    padding-right: calc(1rem + env(safe-area-inset-right, 0px)) !important;
}

section[data-testid="stSidebar"] {
    background-color: var(--ucc-sidebar-bg);
}
[data-testid="stSidebar"] [data-testid="stMarkdown"],
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label {
    color: rgba(255, 255, 255, 0.92);
}

.ucc-inst-header {
    background: var(--ucc-green);
    border-radius: 14px;
    padding: 1.25rem 1.65rem;
    margin-bottom: 1.35rem;
    display: flex;
    flex-direction: row;
    align-items: center;
    gap: 1.35rem;
    flex-wrap: wrap;
    box-sizing: border-box;
}
.ucc-inst-escudo {
    width: 112px;
    max-width: 28vw;
    height: auto;
    flex-shrink: 0;
    display: block;
    object-fit: contain;
    border-radius: 8px;
    background: rgba(255, 255, 255, 0.1);
}
.ucc-inst-banner-text {
    flex: 1 1 240px;
    min-width: 0;
    display: flex;
    flex-direction: column;
    justify-content: center;
}
.header-uccuyo h1.ucc-banner-heading,
.header-uccuyo h2.ucc-banner-heading,
.header-uccuyo h3.ucc-banner-heading {
    color: #ffffff !important;
    margin: 0;
    line-height: 1.2;
    font-family: "Source Sans Pro", ui-sans-serif, system-ui, sans-serif;
}
.header-uccuyo h1.ucc-banner-heading {
    font-size: clamp(1.35rem, 2.8vw, 1.95rem);
    font-weight: 700;
}
.header-uccuyo h2.ucc-banner-heading {
    margin-top: 0.55rem !important;
    font-size: clamp(1rem, 2vw, 1.25rem);
    font-weight: 500;
}
.header-uccuyo h3.ucc-banner-heading {
    margin-top: 0.35rem !important;
    font-size: clamp(0.85rem, 1.4vw, 1rem);
    font-weight: 400;
    color: rgba(255, 255, 255, 0.92) !important;
}

h1:not(.ucc-banner-heading):not(.uc-card-main-title),
h2:not(.ucc-banner-heading),
h3:not(.ucc-banner-heading),
h4 {
    color: var(--ucc-green-dark) !important;
}

.ucc-intro-card {
    background: #ffffff;
    border-radius: 14px;
    padding: 1.75rem 2rem;
    margin-bottom: 1.65rem;
    box-shadow:
        0 8px 28px rgba(0, 0, 0, 0.07),
        0 1px 3px rgba(0, 0, 0, 0.04);
}
.ucc-intro-card h1.uc-card-main-title {
    color: var(--ucc-heading-card) !important;
    margin: 0 0 0.75rem 0 !important;
    font-size: clamp(1.3rem, 2.8vw, 1.85rem);
    font-weight: 700;
    line-height: 1.25;
    font-family: "Source Sans Pro", ui-sans-serif, system-ui, sans-serif;
}
.ucc-intro-card p.uc-card-lead {
    color: var(--ucc-lead-muted) !important;
    margin: 0 !important;
    line-height: 1.6;
    font-size: 1.02rem;
}

p:not(.ucc-banner-heading):not(.uc-card-lead),
label {
    color: var(--ucc-text) !important;
}

[data-testid="stTextInput"] input,
[data-testid="stNumberInput"] input,
[data-testid="stTextArea"] textarea {
    border-radius: 12px !important;
    border: 1px solid rgba(0, 82, 62, 0.22) !important;
    background-color: #ffffff !important;
    color: var(--ucc-text) !important;
    caret-color: var(--ucc-green-dark) !important;
}
[data-baseweb="select"] > div:first-child {
    border-radius: 12px !important;
}

[data-testid="stFileUploader"] {
    background-color: transparent !important;
    border: none !important;
    padding: 0 !important;
}
[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] {
    background-color: #1e1e1e !important;
    border-radius: 12px !important;
    border: 1px solid rgba(255, 255, 255, 0.08) !important;
    padding: 0.85rem 1rem !important;
}
[data-testid="stFileUploaderDropzone"] label,
[data-testid="stFileUploaderDropzone"] span,
[data-testid="stFileUploaderDropzone"] p,
[data-testid="stFileUploaderDropzone"] small,
[data-testid="stFileUploader"] [data-testid="stMarkdownContainer"] p,
[data-testid="stFileUploader"] [data-testid="stMarkdownContainer"] span {
    color: rgba(255, 255, 255, 0.92) !important;
}

[data-testid="stBaseButton-primary"],
[data-testid="stBaseButton-secondary"],
[data-testid="stBaseButton-tertiary"] {
    background-color: var(--ucc-green) !important;
    color: #ffffff !important;
    border-color: transparent !important;
    --text-color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
    font-weight: 600 !important;
}
[data-testid="stBaseButton-primary"]:hover,
[data-testid="stBaseButton-secondary"]:hover,
[data-testid="stBaseButton-tertiary"]:hover {
    background-color: var(--ucc-green-dark) !important;
    border-color: transparent !important;
    color: #ffffff !important;
    --text-color: #ffffff !important;
}
[data-testid="stBaseButton-primary"] p,
[data-testid="stBaseButton-primary"] span,
[data-testid="stBaseButton-secondary"] p,
[data-testid="stBaseButton-secondary"] span,
[data-testid="stBaseButton-tertiary"] p,
[data-testid="stBaseButton-tertiary"] span,
[data-testid="stBaseButton-primary"] div,
[data-testid="stBaseButton-secondary"] div,
[data-testid="stBaseButton-tertiary"] div {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}
[data-testid="stBaseButton-primary"] svg,
[data-testid="stBaseButton-secondary"] svg,
[data-testid="stBaseButton-tertiary"] svg,
[data-testid="stFileUploader"] button svg {
    fill: #ffffff !important;
    color: #ffffff !important;
}

.stButton > button,
[data-testid="stDownloadButton"] button,
[data-testid="stFileUploader"] button {
    background-color: var(--ucc-green) !important;
    color: #ffffff !important;
    border-radius: 8px !important;
    border: none !important;
    font-weight: 600 !important;
    --text-color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}
.stButton > button:hover,
[data-testid="stDownloadButton"] button:hover,
[data-testid="stFileUploader"] button:hover {
    background-color: var(--ucc-green-dark) !important;
    border-color: transparent !important;
}
.stButton > button *,
[data-testid="stDownloadButton"] button *,
[data-testid="stFileUploader"] button * {
    color: #ffffff !important;
    -webkit-text-fill-color: #ffffff !important;
}

div[data-testid="stAlert"] {
    border-radius: 10px;
}

.stSlider label,
[data-testid="stTextInput"] label,
[data-testid="stTextArea"] label,
[data-testid="stFileUploader"] label {
    position: relative;
    padding-left: 1rem;
}
.stSlider label::before,
[data-testid="stTextInput"] label::before,
[data-testid="stTextArea"] label::before,
[data-testid="stFileUploader"] label::before {
    content: "";
    position: absolute;
    left: 0;
    top: 0.45rem;
    width: 9px;
    height: 9px;
    border-radius: 50%;
    background: var(--ucc-accent);
}
</style>
"""

st.set_page_config(page_title="Valorador de CVar CLEAN - UCCuyo (TXT)", layout="wide")

st.markdown(_UCCI_GLOBAL_CSS, unsafe_allow_html=True)

_inst_header_html = f"""
<div class="ucc-inst-header header-uccuyo">
<img class="ucc-inst-escudo" src="{_escudo_src_for_banner()}" alt="Universidad Católica de Cuyo" />
<div class="ucc-inst-banner-text">
<h1 class="ucc-banner-heading">Universidad Católica de Cuyo</h1>
<h2 class="ucc-banner-heading">Secretaría de Investigación</h2>
<h3 class="ucc-banner-heading">Consejo de Investigación</h3>
</div>
</div>
"""
st.markdown(_inst_header_html, unsafe_allow_html=True)

st.markdown(
    """
<div class="ucc-intro-card">
<h1 class="uc-card-main-title">Universidad Católica de Cuyo — Valorador de CVar (TXT limpio)</h1>
<p class="uc-card-lead">Entrada: *_CVAR_CLEAN.txt (salida del Normalizador). Exporta Excel y Word + categoría automática.
(STRICT: no puntúa títulos sin finalización explícita)</p>
</div>
""",
    unsafe_allow_html=True,
)

criteria = load_criteria("criteria.json")

DEBUG = st.checkbox("Debug (mostrar vista previa y coincidencias)", value=False)

uploaded = st.file_uploader("Cargar CVar limpio (.txt)", type=["txt"])
if not uploaded:
    st.info("Subí un archivo TXT limpio para iniciar la valoración.")
    st.stop()

raw = uploaded.read()
text = raw.decode("utf-8", errors="ignore")

st.success(f"Archivo cargado: {uploaded.name}")

preview_lines = 200
lines = text.splitlines()
preview = "\n".join(lines[:preview_lines])

if DEBUG:
    with st.expander(f"Vista previa (primeras {preview_lines} líneas)"):
        st.code(preview if preview.strip() else "(archivo vacío)", language="text")

with st.spinner("Calculando puntajes con criteria.json..."):
    item_results, section_totals, total, category, categorias = score_text(text, criteria)

desc_cat = categorias.get(category, {}).get("descripcion", "")

st.markdown("---")
st.subheader("Puntaje total y categoría")
st.metric("Total acumulado", f"{total:.1f}")
st.metric("Categoría alcanzada", f"Categoría {category}")
if desc_cat:
    st.info(f"Descripción de la categoría: {desc_cat}")

rows = []
for r in item_results:
    rows.append(
        {
            "Sección": r.section,
            "Ítem": r.item,
            "Ocurrencias": r.count,
            "Unit points": r.unit_points,
            "Puntaje bruto": r.raw_points,
            "Tope ítem": r.item_max_points,
            "Puntaje (tope aplicado)": r.capped_item_points,
            "Evidencia (1er match)": r.evidence,
        }
    )

df_items = pd.DataFrame(rows)

st.markdown("---")
st.subheader("Detalle por sección")

for section_name, cfg in criteria.get("sections", {}).items():
    st.markdown(f"### {section_name}")
    df_sec = df_items[df_items["Sección"] == section_name].copy()
    df_sec = df_sec.sort_values(["Puntaje (tope aplicado)", "Ocurrencias"], ascending=False)

    st.dataframe(
        df_sec[
            ["Ítem", "Ocurrencias", "Puntaje (tope aplicado)", "Tope ítem", "Evidencia (1er match)"]
        ],
        use_container_width=True,
        hide_index=True,
    )

    sec_max = cfg.get("max_points", 0)
    sec_sub = section_totals.get(section_name, 0.0)
    st.info(f"Subtotal {section_name}: {sec_sub:.1f} / máx {sec_max}")

st.markdown("---")
st.subheader("Totales por sección (tope de sección aplicado)")
df_sec_tot = pd.DataFrame([{"Sección": k, "Subtotal": v} for k, v in section_totals.items()])
df_sec_tot = df_sec_tot.sort_values("Subtotal", ascending=False)
st.dataframe(df_sec_tot, use_container_width=True, hide_index=True)

st.markdown("---")
st.subheader("Exportar resultados")

excel_out = io.BytesIO()
with pd.ExcelWriter(excel_out, engine="xlsxwriter") as writer:
    for section_name in criteria.get("sections", {}).keys():
        df_s = df_items[df_items["Sección"] == section_name].copy()
        df_s.to_excel(writer, sheet_name=section_name[:31], index=False)

    resumen = df_sec_tot.copy()
    resumen.loc[len(resumen)] = ["TOTAL", total]
    resumen.loc[len(resumen)] = ["CATEGORÍA", f"Categoría {category}"]
    resumen.to_excel(writer, sheet_name="RESUMEN", index=False)

excel_out.seek(0)

st.download_button(
    "Descargar Excel",
    data=excel_out.getvalue(),
    file_name=uploaded.name.replace(".txt", "") + "__PUNTAJE.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    type="primary",
    use_container_width=True,
)


def export_word(df_items_local, df_sec_tot_local, total_pts, cat, cat_desc, filename):
    doc = DocxDocument()
    p = doc.add_paragraph("Universidad Católica de Cuyo — Secretaría de Investigación")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Informe de valoración de CVar (TXT limpio) — STRICT").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"Archivo evaluado: {filename}")
    doc.add_paragraph(f"Puntaje total: {total_pts:.1f}")
    doc.add_paragraph(f"Categoría alcanzada: Categoría {cat}")
    if cat_desc:
        doc.add_paragraph(cat_desc)

    doc.add_paragraph("")
    doc.add_heading("Totales por sección", level=2)
    for _, row in df_sec_tot_local.iterrows():
        doc.add_paragraph(f"- {row['Sección']}: {float(row['Subtotal']):.1f}")

    for section_name in criteria.get("sections", {}).keys():
        doc.add_heading(section_name, level=2)
        df_s = df_items_local[df_items_local["Sección"] == section_name].copy()

        cols = ["Ítem", "Ocurrencias", "Puntaje (tope aplicado)", "Tope ítem"]
        if DEBUG:
            cols.append("Evidencia (1er match)")

        if df_s.empty:
            doc.add_paragraph("Sin ítems detectados.")
        else:
            tbl = doc.add_table(rows=1, cols=len(cols))
            hdr = tbl.rows[0].cells
            for i, c in enumerate(cols):
                hdr[i].text = c

            for _, r in df_s.iterrows():
                cells = tbl.add_row().cells
                for i, c in enumerate(cols):
                    cells[i].text = str(r.get(c, ""))

        doc.add_paragraph(
            f"Subtotal sección: {df_sec_tot_local[df_sec_tot_local['Sección'] == section_name]['Subtotal'].values[0]:.1f}"
            if (df_sec_tot_local["Sección"] == section_name).any()
            else ""
        )

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


st.download_button(
    "Descargar informe Word",
    data=export_word(df_items, df_sec_tot, total, category, desc_cat, uploaded.name),
    file_name=uploaded.name.replace(".txt", "") + "__INFORME.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    type="primary",
    use_container_width=True,
)
